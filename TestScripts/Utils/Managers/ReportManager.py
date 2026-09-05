# Import Required Libraries

try:
    from docx import Document
    from docx.shared import Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    Document = None
    Inches = None
    OxmlElement = None
    qn = None

try:
    from docx2pdf import convert as convert_to_pdf
except ImportError:
    convert_to_pdf = None

try:
    import pythoncom
    from win32com.client import Dispatch
except ImportError:
    pythoncom = None
    Dispatch = None

from TestScripts.Utils.Managers.CommonFunctions import *
from TestScripts.Utils.Managers.GlosssaryManager import *


class ReportManager:
    """Generate a simple DOCX report for all executed tests."""

    m_sReportFile = os.getcwd() + "//TestResults//WinTestReport.docx"
    m_sPdfReportFile = os.getcwd() + "//TestResults//WinTestReport.pdf"
    m_sFallbackReportFile = os.getcwd() + "//TestResults//WinTestReport.txt"
    m_CF = CommonFunctions()

    def __init__(self):
        logging.info("ReportManager::ReportManager()")
        self.m_sReportFile = os.getcwd() + "//TestResults//WinTestReport.docx"
        self.m_sPdfReportFile = os.getcwd() + "//TestResults//WinTestReport.pdf"
        self.m_sFallbackReportFile = os.getcwd() + "//TestResults//WinTestReport.txt"
        self.m_Document = Document() if Document is not None else None

    def WriteReport(self, testResultsList):
        logging.info("ReportManager::WriteReport()")
        if Document is None:
            logging.warning(
                "python-docx is not installed. Generating a text report instead of a Word document."
            )
            self._write_text_report(testResultsList)
            return

        self.m_Document.add_heading("Windows Application Test Results", level=0)
        self.m_Document.add_paragraph("Generated on: " + self.m_CF.GetCurrentTime())
        self.m_Document.add_page_break()

        self.m_Document.add_heading("Index", level=1)
        self._add_table_of_contents()
        self.m_Document.add_page_break()

        self.m_Document.add_heading("1. Test Summary", level=1)
        table = self.m_Document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = table.rows[0].cells
        headers[0].text = "ID"
        headers[1].text = "Description"
        headers[2].text = "Status"
        headers[3].text = "Start Time"
        headers[4].text = "End Time"

        for test in testResultsList:
            row = table.add_row().cells
            row[0].text = test.get_TestID()
            row[1].text = test.get_TestDescription()
            row[2].text = test.get_TestResult()
            row[3].text = test.get_TestStartTime()
            row[4].text = test.get_TestEndTime()

        for index, test in enumerate(testResultsList, start=2):
            self.m_Document.add_page_break()
            self.m_Document.add_heading(str(index) + ". " + test.get_TestID(), level=1)
            self.m_Document.add_paragraph("Description: " + test.get_TestDescription())
            self.m_Document.add_paragraph("Status: " + test.get_TestResult())

            if self.m_CF.GetActualResultsInReport() == 1:
                self.m_Document.add_heading("Actual Results", level=2)
                self._append_text_files(
                    os.getcwd() + "\\TestResults\\" + test.get_ModuleName() + "\\" + test.get_TestID()
                )

            if self.m_CF.GetExpectedResultsInReport() == 1:
                self.m_Document.add_heading("Expected Results", level=2)
                self._append_text_files(
                    os.getcwd() + "\\ExpectedResults\\" + test.get_ModuleName() + "\\" + test.get_TestID()
                )

            if self.m_CF.GetScreenshotsInReport() == 1:
                screenshotFolder = os.getcwd() + "\\TestResults\\" + test.get_ModuleName() + "\\" + test.get_TestID()
                jpgFiles = sorted(
                    [name for name in os.listdir(screenshotFolder) if name.lower().endswith(".jpg")]
                )
                if jpgFiles:
                    self.m_Document.add_heading("Screenshots", level=2)
                    for fileName in jpgFiles:
                        self.m_Document.add_paragraph(fileName)
                        self.m_Document.add_picture(screenshotFolder + "\\" + fileName, width=Inches(5))

        if self.m_CF.GetGlossaryInReport() == 1:
            self.m_Document.add_page_break()
            self.m_Document.add_heading("Glossary", level=1)
            glossary = GlossaryManager().GetGlossary()
            table = self.m_Document.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Term"
            table.rows[0].cells[1].text = "Explanation"
            for key, value in glossary.items():
                row = table.add_row().cells
                row[0].text = str(key)
                row[1].text = str(value)

        self.m_Document.save(self.m_sReportFile)
        self._update_table_of_contents()
        self._generate_pdf_report_if_enabled()

    def _write_text_report(self, testResultsList):
        with open(self.m_sFallbackReportFile, "w", encoding="utf-8") as report:
            report.write("Windows Application Test Results\n")
            report.write("Generated on: " + self.m_CF.GetCurrentTime() + "\n\n")
            report.write("Test Summary\n")
            report.write("=" * 60 + "\n")

            for test in testResultsList:
                report.write("ID: " + test.get_TestID() + "\n")
                report.write("Description: " + test.get_TestDescription() + "\n")
                report.write("Status: " + test.get_TestResult() + "\n")
                report.write("Start Time: " + test.get_TestStartTime() + "\n")
                report.write("End Time: " + test.get_TestEndTime() + "\n")
                report.write("-" * 60 + "\n")

                actualFile = os.getcwd() + "\\TestResults\\" + test.get_ModuleName() + "\\" + test.get_TestID() + "\\TestResult.txt"
                expectedFile = os.getcwd() + "\\ExpectedResults\\" + test.get_ModuleName() + "\\" + test.get_TestID() + "\\TestResult.txt"

                if os.path.isfile(actualFile):
                    report.write("Actual Result:\n")
                    with open(actualFile, encoding="utf-8") as handle:
                        report.write(handle.read() + "\n")

                if os.path.isfile(expectedFile):
                    report.write("Expected Result:\n")
                    with open(expectedFile, encoding="utf-8") as handle:
                        report.write(handle.read() + "\n")

                report.write("\n")

            if self.m_CF.GetGlossaryInReport() == 1:
                report.write("Glossary\n")
                report.write("=" * 60 + "\n")
                glossary = GlossaryManager().GetGlossary()
                for key, value in glossary.items():
                    report.write(str(key) + ": " + str(value) + "\n")

    def _append_text_files(self, folderPath):
        if not os.path.isdir(folderPath):
            return

        for fileName in sorted(os.listdir(folderPath)):
            if fileName.lower().endswith(".txt"):
                self.m_Document.add_paragraph(fileName)
                with open(folderPath + "\\" + fileName, encoding="utf-8") as handle:
                    self.m_Document.add_paragraph(handle.read())

    def _add_table_of_contents(self):
        if OxmlElement is None or qn is None:
            self.m_Document.add_paragraph("Table of contents is unavailable because python-docx TOC support is not loaded.")
            return

        paragraph = self.m_Document.add_paragraph()
        run = paragraph.add_run()

        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")

        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = 'TOC \\o "1-3" \\h \\z \\u'

        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")

        placeholder = OxmlElement("w:t")
        placeholder.text = "Right-click to update field."
        separate.append(placeholder)

        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")

        run._r.append(begin)
        run._r.append(instruction)
        run._r.append(separate)
        run._r.append(end)

    def _update_table_of_contents(self):
        if Dispatch is None or pythoncom is None:
            logging.warning(
                "Table of contents was added to the document, but pywin32 is not installed so it could not be refreshed automatically."
            )
            return

        try:
            pythoncom.CoInitialize()
            word = Dispatch("Word.Application")
            word.Visible = False
            document = word.Documents.Open(os.path.abspath(self.m_sReportFile))

            if document.TablesOfContents.Count >= 1:
                document.TablesOfContents(1).Update()

            document.Save()
            document.Close(False)
            word.Quit()
            logging.info("ReportManager::_update_table_of_contents() - TOC updated successfully")
        except Exception:
            logging.exception(
                "Table of contents was added, but automatic refresh failed. Open the DOCX in Word and update the TOC manually if needed."
            )
        finally:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    def _generate_pdf_report_if_enabled(self):
        if self.m_CF.GetPDFGenerationInReport() != 1:
            return

        if Document is None:
            logging.warning(
                "PDF report requested, but python-docx is unavailable so the DOCX report was not created."
            )
            return

        if convert_to_pdf is None:
            logging.warning(
                "PDF report requested, but docx2pdf is not installed. Install docx2pdf to enable PDF generation."
            )
            return

        try:
            convert_to_pdf(self.m_sReportFile, self.m_sPdfReportFile)
            logging.info("ReportManager::WriteReport() - Generated PDF report: %s", self.m_sPdfReportFile)
        except Exception:
            logging.exception(
                "PDF report requested, but conversion failed. Microsoft Word may be required for docx2pdf on Windows."
            )
